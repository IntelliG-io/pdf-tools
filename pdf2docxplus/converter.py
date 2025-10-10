"""Conversion engine for pdf2docxplus."""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from pdf2docx import Converter

from .exceptions import ConversionError
from .metadata import PDFMetadata, apply_metadata_to_docx, extract_metadata
from .utils import ensure_output_directory, time_block, to_path
from .validators import validate_conversion, validate_pdf

LOGGER = logging.getLogger(__name__)


@dataclass(frozen=True)
class ConversionOptions:
    """Options controlling PDF to DOCX conversion."""

    preserve_formatting: bool = True
    include_metadata: bool = True
    optimize: bool = False


def _apply_layout_enhancements(docx_path: Path) -> None:
    """Perform post-processing on the DOCX to enhance layout fidelity."""
    document = Document(str(docx_path))
    for table in document.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
    document.save(str(docx_path))


def convert_pdf_to_docx(
    input_path: str | Path,
    output_path: str | Path,
    preserve_formatting: bool = True,
    include_metadata: bool = True,
    optimize: bool | None = None,
) -> Path:
    """Convert a PDF file to DOCX with high fidelity."""
    options = ConversionOptions(
        preserve_formatting=preserve_formatting,
        include_metadata=include_metadata,
        optimize=bool(optimize) if optimize is not None else False,
    )

    source = to_path(input_path)
    destination = to_path(output_path)
    ensure_output_directory(destination)

    LOGGER.info("Starting conversion: %s -> %s", source, destination)
    validate_pdf(source)

    extracted_metadata: Optional[PDFMetadata] = None
    if options.include_metadata:
        try:
            extracted_metadata = extract_metadata(source)
        except Exception as exc:  # pragma: no cover - handled via MetadataError
            LOGGER.warning("Metadata extraction failed: %s", exc)

    layout_mode = "exact" if options.preserve_formatting else "normal"
    image_mode = "strict" if options.preserve_formatting else "compress"

    try:
        with time_block(LOGGER, "PDF to DOCX conversion"):
            with Converter(str(source)) as converter:
                converter.convert(
                    str(destination),
                    start=0,
                    end=None,
                    layout_mode=layout_mode,
                    image_mode=image_mode,
                    tables="detect",
                    text="preserve",
                    progress_bar=False,
                )
    except Exception as exc:  # pragma: no cover
        raise ConversionError(f"Conversion failed for {source}") from exc

    if options.preserve_formatting:
        try:
            _apply_layout_enhancements(destination)
        except Exception as exc:  # pragma: no cover
            LOGGER.debug("Layout enhancements failed: %s", exc)

    if extracted_metadata and options.include_metadata:
        try:
            apply_metadata_to_docx(destination, extracted_metadata)
        except Exception as exc:  # pragma: no cover
            LOGGER.warning("Failed to apply metadata: %s", exc)

    validate_conversion(destination, extracted_metadata if options.include_metadata else None)
    LOGGER.info("Conversion completed: %s", destination)
    return destination
