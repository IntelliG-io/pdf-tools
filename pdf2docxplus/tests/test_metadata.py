from __future__ import annotations

from datetime import timezone

from docx import Document

from pdf2docxplus.metadata import PDFMetadata, apply_metadata_to_docx, extract_metadata


def test_extract_metadata(sample_pdf):
    metadata = extract_metadata(sample_pdf)
    assert metadata.title == "Test Document"
    assert metadata.author == "Jane Doe"
    assert metadata.subject == "Testing"
    assert metadata.keywords == "pdf,docx"
    assert metadata.creator == "pytest"
    assert metadata.producer == "pdf2docxplus"
    assert metadata.additional_properties == {"Custom": "Value"}
    assert metadata.creation_date.tzinfo is not None
    assert metadata.modification_date.tzinfo is not None


def test_apply_metadata_to_docx(tmp_path, sample_pdf):
    docx_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("Hello")
    document.save(docx_path)

    metadata = extract_metadata(sample_pdf)
    apply_metadata_to_docx(docx_path, metadata)

    reloaded = Document(docx_path)
    props = reloaded.core_properties
    assert props.title == metadata.title
    assert props.author == metadata.author
    assert props.subject == metadata.subject
    assert props.keywords == metadata.keywords
    assert props.created == metadata.creation_date.astimezone(timezone.utc)
    assert props.modified == metadata.modification_date.astimezone(timezone.utc)
